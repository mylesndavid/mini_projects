from docx import Document

document = Document("Mockingbirds.docx")
paragraph = document.paragraphs[0]
print(paragraph.text)

passage = paragraph.text.split()

para = document.add_paragraph()

for word in passage:
    halfWord = int(len(word)/2)
    for letters in word:
        firstHalf = word[0:halfWord]
        restOfWord = word.replace(firstHalf,'')
        if halfWord < 1:
            firstHalf = word
            restOfWord = ""
    bold_letters = para.add_run(firstHalf)
    bold_letters.bold = True
    para.add_run(restOfWord)
    para.add_run(" ")

document.save("Mockingbirds.docx")